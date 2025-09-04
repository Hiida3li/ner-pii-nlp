"""
Document Processing Module
Handles text extraction from various document formats with PII detection integration
"""

import os
import io
import logging
from typing import Dict, List, Optional, Tuple, Any
from pathlib import Path
import hashlib
from datetime import datetime
import re

# Document processing libraries
import PyPDF2
from docx import Document
import pandas as pd
import openpyxl
from openpyxl import load_workbook

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Document processing class that extracts text from various formats
    and integrates with the PII detection system
    """
    
    SUPPORTED_FORMATS = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.doc': 'application/msword',
        '.txt': 'text/plain',
        '.md': 'text/markdown',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.xls': 'application/vnd.ms-excel',
        '.csv': 'text/csv'
    }
    
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    MAX_TEXT_LENGTH = 100000  # 100k characters
    
    def __init__(self):
        """Initialize the document processor"""
        self.temp_dir = Path("temp_uploads")
        self.temp_dir.mkdir(exist_ok=True)
        
    def is_supported_format(self, filename: str) -> bool:
        """Check if file format is supported"""
        file_ext = Path(filename).suffix.lower()
        return file_ext in self.SUPPORTED_FORMATS
    
    def get_file_info(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract file metadata"""
        file_hash = hashlib.md5(file_content).hexdigest()
        file_size = len(file_content)
        file_ext = Path(filename).suffix.lower()
        
        return {
            'filename': filename,
            'size': file_size,
            'hash': file_hash,
            'extension': file_ext,
            'mime_type': self.SUPPORTED_FORMATS.get(file_ext, 'unknown'),
            'uploaded_at': datetime.now().isoformat()
        }
    
    async def process_document(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Main document processing function
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            # Validate file
            if len(file_content) > self.MAX_FILE_SIZE:
                raise ValueError(f"File too large. Maximum size: {self.MAX_FILE_SIZE / (1024*1024):.1f}MB")
            
            if not self.is_supported_format(filename):
                raise ValueError(f"Unsupported file format. Supported: {', '.join(self.SUPPORTED_FORMATS.keys())}")
            
            # Get file info
            file_info = self.get_file_info(file_content, filename)
            
            # Extract text based on file type
            file_ext = Path(filename).suffix.lower()
            
            if file_ext == '.pdf':
                extracted_text = self._extract_from_pdf(file_content)
            elif file_ext in ['.docx', '.doc']:
                extracted_text = self._extract_from_docx(file_content)
            elif file_ext in ['.txt', '.md']:
                extracted_text = self._extract_from_text(file_content)
            elif file_ext in ['.xlsx', '.xls']:
                extracted_text = self._extract_from_excel(file_content)
            elif file_ext == '.csv':
                extracted_text = self._extract_from_csv(file_content)
            else:
                raise ValueError(f"No processor available for {file_ext}")
            
            # Clean and validate extracted text
            cleaned_text = self._clean_extracted_text(extracted_text)
            
            if len(cleaned_text) > self.MAX_TEXT_LENGTH:
                cleaned_text = cleaned_text[:self.MAX_TEXT_LENGTH] + "\n\n[Text truncated due to length limit]"
            
            # Prepare result
            result = {
                'success': True,
                'text': cleaned_text,
                'file_info': file_info,
                'text_length': len(cleaned_text),
                'word_count': len(cleaned_text.split()),
                'processing_time': datetime.now().isoformat()
            }
            
            logger.info(f"Successfully processed {filename}: {len(cleaned_text)} characters extracted")
            return result
            
        except Exception as e:
            logger.error(f"Error processing document {filename}: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'file_info': self.get_file_info(file_content, filename) if file_content else None
            }
    
    def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF files"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"[Page {page_num + 1}]\n{page_text}\n")
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num + 1}: {e}")
                    text_parts.append(f"[Page {page_num + 1} - Extraction Error]\n")
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise ValueError("Failed to extract text from PDF file")
    
    def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX files"""
        try:
            docx_file = io.BytesIO(file_content)
            document = Document(docx_file)
            
            text_parts = []
            
            # Extract paragraphs
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in document.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(' | '.join(row_text))
                
                if table_text:
                    text_parts.append('\n[Table]\n' + '\n'.join(table_text) + '\n[End Table]\n')
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise ValueError("Failed to extract text from DOCX file")
    
    def _extract_from_text(self, file_content: bytes) -> str:
        """Extract text from plain text files"""
        try:
            # Try different encodings
            for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Could not decode text file")
            
        except Exception as e:
            logger.error(f"Text extraction error: {e}")
            raise ValueError("Failed to extract text from file")
    
    def _extract_from_excel(self, file_content: bytes) -> str:
        """Extract text from Excel files"""
        try:
            excel_file = io.BytesIO(file_content)
            workbook = load_workbook(excel_file, read_only=True, data_only=True)
            
            text_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"\n[Sheet: {sheet_name}]\n")
                
                # Extract data from cells
                sheet_data = []
                for row in sheet.iter_rows(values_only=True):
                    if any(cell for cell in row if cell is not None):
                        row_text = []
                        for cell in row:
                            if cell is not None:
                                row_text.append(str(cell))
                        if row_text:
                            sheet_data.append(' | '.join(row_text))
                
                if sheet_data:
                    text_parts.extend(sheet_data)
                text_parts.append(f"[End Sheet: {sheet_name}]\n")
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"Excel extraction error: {e}")
            raise ValueError("Failed to extract text from Excel file")
    
    def _extract_from_csv(self, file_content: bytes) -> str:
        """Extract text from CSV files"""
        try:
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    csv_text = file_content.decode(encoding)
                    # Parse CSV using pandas for better handling
                    df = pd.read_csv(io.StringIO(csv_text))
                    
                    # Convert to formatted text
                    text_parts = ["[CSV Data]\n"]
                    
                    # Add column headers
                    headers = ' | '.join(df.columns.astype(str))
                    text_parts.append(headers)
                    text_parts.append('-' * len(headers))
                    
                    # Add data rows
                    for _, row in df.iterrows():
                        row_text = ' | '.join(row.astype(str))
                        text_parts.append(row_text)
                    
                    text_parts.append("\n[End CSV Data]")
                    return '\n'.join(text_parts)
                    
                except (UnicodeDecodeError, pd.errors.ParserError):
                    continue
            
            raise ValueError("Could not parse CSV file")
            
        except Exception as e:
            logger.error(f"CSV extraction error: {e}")
            raise ValueError("Failed to extract text from CSV file")
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Multiple empty lines
        text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs
        
        # Remove control characters except newlines and tabs
        text = ''.join(char for char in text if char.isprintable() or char in '\n\t')
        
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        return text.strip()
    
    def get_supported_formats(self) -> Dict[str, str]:
        """Return supported file formats"""
        return self.SUPPORTED_FORMATS.copy()
    
    def cleanup_temp_files(self, max_age_hours: int = 1):
        """Clean up old temporary files"""
        try:
            if not self.temp_dir.exists():
                return
            
            max_age_seconds = max_age_hours * 3600
            current_time = datetime.now().timestamp()
            
            for file_path in self.temp_dir.iterdir():
                if file_path.is_file():
                    file_age = current_time - file_path.stat().st_mtime
                    if file_age > max_age_seconds:
                        file_path.unlink()
                        logger.info(f"Cleaned up temp file: {file_path.name}")
                        
        except Exception as e:
            logger.error(f"Error cleaning temp files: {e}")